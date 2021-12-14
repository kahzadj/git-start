from docx import Document
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

import docx
import docx2txt

from tkinter import *
import tkinter.filedialog
import os
import sys
#from googletrans import Translator
import webbrowser


#-------read sorce file for translate---------
root = Tk()
root.withdraw()
currdir = os.getcwd()
document = Document()
document5 = Document()
k = 0
tempfile1 = tkinter.filedialog.askopenfilename(title='Please select a file')
point1 = tempfile1[:len(tempfile1)-5]
if not os.path.exists(point1):
    os.makedirs(point1)
point = tempfile1[:len(tempfile1)-5]+'_images'
if not os.path.exists(point):
    os.makedirs(point)
document1 = Document(tempfile1)

#---------------------------------------


for para in document1.paragraphs:
    paragraph = document.add_paragraph(para.text)

document.save(tempfile1[:len(tempfile1)-5]+'_Paragraphs.docx')
document2 = Document(tempfile1[:len(tempfile1)-5]+'_Paragraphs.docx')
webbrowser.open('https://translate.google.com/?hl=en&tab=rT0&sl=en&tl=fa&op=docs')

document5.save(tempfile1[:len(tempfile1)-5]+'_Translat.docx')
tempfile4 = tkinter.filedialog.askopenfilename(title='Please select a file')
document3 = Document(tempfile4)

#---------Extract images from the original file------
rels = []
docx2txt.process(tempfile1, point)
doc = docx.Document(tempfile1)

for r in doc.part.rels.values():
    if isinstance(r._target, docx.parts.image.ImagePart):
        rels.append(os.path.basename(r._target.partname))
print (rels)
i = 0
for paragraph in doc.paragraphs:
    if paragraph.text == document2.paragraphs[i].text:
        paragraph.text = document3.paragraphs[i].text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    i += 1
doc.save(tempfile1[:len(tempfile1)-5]+'_result.docx')
