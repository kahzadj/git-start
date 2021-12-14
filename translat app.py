from __future__ import (absolute_import, division, print_function, unicode_literals)
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION
import docx
import docx2txt
from docx.shared import Inches


from tkinter import *
from PIL import ImageTk, Image
import tkinter.filedialog
import glob
import os
import io
import sys
from googletrans import Translator
from bidi.algorithm import get_display
from arabic_reshaper import reshape
import time
import webbrowser


#-------read sorce file for translate---------
root = Tk()
root.withdraw()
currdir = os.getcwd()
document = Document()
document5 = Document()
k = 0
tempfile1 = tkinter.filedialog.askopenfilename(title='Please select a file')
point1 = tempfile1[:len(tempfile1)-5]
if not os.path.exists(point1):
    os.makedirs(point1)
point = tempfile1[:len(tempfile1)-5]+'_images'
if not os.path.exists(point):
    os.makedirs(point)
document1 = Document(tempfile1)

#---------------------------------------


for para in document1.paragraphs:
    paragraph = document.add_paragraph(para.text)

document.save(tempfile1[:len(tempfile1)-5]+'_Paragraphs.docx')
document2 = Document(tempfile1[:len(tempfile1)-5]+'_Paragraphs.docx')
webbrowser.open('https://translate.google.com/?hl=en&tab=rT0&sl=en&tl=fa&op=docs')


document5.save(tempfile1[:len(tempfile1)-5]+'_Translat.docx')

tempfile4 = tkinter.filedialog.askopenfilename(title='Please select a file')
document3 = Document(tempfile1[:len(tempfile1)-5]+'_Translat.docx')

#---------Extract images from the original file------
rels = []
docx2txt.process(tempfile1, point)
doc = docx.Document(tempfile1)
for r in doc.part.rels.values():
    if isinstance(r._target, docx.parts.image.ImagePart):
        rels.append(os.path.basename(r._target.partname))

#-----Create a function for navigating the original text items------
def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

#-------Read the text based on the items in the original text-------
i = 0
c = 0

for block in iter_block_items(document1):
    c += 1
    if isinstance(block, Paragraph):        
        if block.text == document2.paragraphs[i].text:
            block.text=block.text.replace(block.text,document3.paragraphs[i].text)
            i += 1

print (c)
document1.save(tempfile1[:len(tempfile1)-5]+'_result.docx')